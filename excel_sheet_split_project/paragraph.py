import os
import re
import pandas as pd
from docx import Document

def identify_bold_letters(doc_path):
    doc = Document(doc_path)
    bold_text = []
    current_bold = ""
    
    for paragraph in doc.paragraphs:
        
        for run in paragraph.runs:
            if run.bold:
                current_bold += "\n" + run.text
            elif current_bold:
                bold_text.append(current_bold.strip())
                current_bold = ""
    if current_bold:
        bold_text.append(current_bold)
    return bold_text

def split_with_first_occurrence(string, delimiters):
    result = [string]
    for delimiter in delimiters:
        new_result = []
        for item in result:
            parts = item.split(delimiter, 1)  
            new_result.extend(parts)
        result = new_result
    return [res for res in result if res != '']

def extract_data_from_documents(documents_dir):
    titles = []
    titles.extend(identify_bold_letters(documents_dir))
    doc = Document(documents_dir)
    breakpoint()
    doc_li = [i.text for i in doc.paragraphs]
    text = '\n'.join(doc_li)
    topics = split_with_first_occurrence(text, titles)

    pattern = r'\n[\s\w()]+:'
    headers = []
    dicts = []
    
    for i, topic in enumerate(topics):
        sub_top = [sub.replace('\n', '').strip(',') for sub in re.split(pattern, topic) if sub != '']
        sub_head = [head.replace('\n', '').replace(':', '') for head in re.findall(pattern, str(topic))]
        headers.extend(sub_head)
        breakpoint()
        items = {}
        items['Product_name'] = titles[i]
        
        for i in range(len(sub_head)):
            items[sub_head[i]] = sub_top[i] #to take it as key value pair the headers and the content
        
        dicts.append(items)
    
    return dicts

def write_dicts_to_excel(dicts, output_dir=None, file_name='output_final_data.xlsx'):
    fieldnames = ['Product_name', 'Short Description', 'Our View', 'Rating', 'Popularity', 'Location', 'Recommended for','Good (Recommended) for', 'Not Recommended for', 'Do Not Miss', 'Tips', 'Timings', 'Duration', 'Tickets and Pricing', 'Inclusions and Exclusions', 'Things to do nearby', 'Guide Options', 'Public Transport', 'Wheelchair Accessibility', 'Website']
    
    df = pd.DataFrame(dicts)
    df = df[fieldnames]  
    df['Recommended for'] = df['Recommended for'].fillna('') + ' ' + df['Good (Recommended) for'].fillna('')
    df.drop(columns=['Good (Recommended) for'], inplace=True)

    if output_dir is None:
        output_dir = os.getcwd()

    output_path = os.path.join(output_dir, file_name)
    os.makedirs(output_dir, exist_ok=True)
    df.to_excel(output_path, index=False)

curr_dir = os.path.join(os.path.abspath('files')) 
list_file = os.listdir(curr_dir)
# breakpoint()
file_paths = [os.path.join(curr_dir, f) for f in list_file if os.path.isfile(os.path.join(curr_dir, f))]

all_data_dicts = []

for file_path in file_paths:
    data_dicts = extract_data_from_documents(file_path)
    all_data_dicts.extend(data_dicts)

write_dicts_to_excel(all_data_dicts)

