import os
import re
import logging
import pandas as pd
from docx import Document


logging.basicConfig(filename='app.log', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def read_document_and_extract_data(doc_path):
    try:
        doc = Document(doc_path)
        bold_text = []
        current_bold = ""
        headers = []
        dicts = []

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.bold:
                    current_bold += "\n" + run.text
                elif current_bold:  
                    bold_text.append(current_bold.strip())
                    current_bold = ""
        
        if current_bold:
            bold_text.append(current_bold)

        doc_text = '\n'.join([i.text for i in doc.paragraphs])
        topics = split_with_first_occurrence(doc_text, bold_text)

        pattern = r"\n\s*[\s\w()'?]+:\s*"

        for i, topic in enumerate(topics):
            if not re.search(r'\\n\\n\s*Rating:\d\\n\\n | \n\nRating:\d\n\n',topic):
                topic = topic.replace('\n\nRating:\n\n','\n\nRating:0\n\n').replace('\n\n Rating:\n\n','\n\nRating:0\n\n')
            elif not re.search(r'\n\n\s*Rating:\d\n\n',topic):

                topic = topic.replace('\n\n\s*Rating:\n\n','\n\n\s*Rating:0\n\n')
            sub_top = [sub.replace('\n', '\n').strip(',') for sub in re.split(pattern, topic) if sub != '']
            sub_head = [head.replace('\n', '').replace(':', '') for head in re.findall(pattern, str(topic))]
            sub_head = [sub.strip() for sub in sub_head]
            sub_top = [sub.strip() for sub in sub_top]
            headers.extend(sub_head)

            items = {}
            # if i < len(bold_text):  
            items['Product_name'] = bold_text[i]
            for i in range(len(sub_head)):
                if i < len(sub_top): 
                    items[sub_head[i]] = sub_top[i]
                else:
                    items[sub_head[i]] = ''  
            dicts.append(items)

        return dicts, headers
    except Exception as e:
        logging.error(f"Error in read_document_and_extract_data function: {e}")
        return [], []

def split_with_first_occurrence(string, delimiters):
    try:
        result = [string]
        for delimiter in delimiters:
            new_result = []
            for item in result:
                parts = item.split(delimiter, 1)  
                new_result.extend(parts)
            result = new_result
        return [res for res in result if res != '']
    except Exception as e:
        logging.error(f"Error in split_with_first_occurrence function: {e}")
        return []

def write_dicts_to_excel(dicts, output_dir=None, file_name='output_final_data_6.xlsx'):
    try:
        df = pd.DataFrame(dicts)
        
        if 'Good (Recommended) for' in df.columns:
            df['Recommended for'] = df['Recommended for'].fillna('') + ' ' + df['Good (Recommended) for'].fillna('')
            df.drop(columns=['Good (Recommended) for'], inplace=True)

        actual_columns = df.columns.tolist()

        df = df[actual_columns]
        
        if 'What to eat' in df.columns and 'What to Eat' in df.columns:
            df['What to Eat'] = df['What to Eat'].fillna('') + ' ' + df['What to eat'].fillna('')
            df.drop(columns=['What to eat'], inplace=True)
        df = df.loc[:, ~df.columns.duplicated()]

        if 'Do Not Miss' in df.columns and "Don't Miss" in df.columns:
            df["Don't Miss"] = df["Don't Miss"].fillna('') + ' ' + df['Do Not Miss'].fillna('')
            df.drop(columns=['Do Not Miss'], inplace=True)
        df = df.loc[:, ~df.columns.duplicated()] 

        if 'Where to Eat?' in df.columns and 'Where to eat?' in df.columns:
            df['Where to Eat?'] = df['Where to Eat?'].fillna('') + ' ' + df['Where to eat?'].fillna('')
            df.drop(columns=['Where to eat?'], inplace=True)
        df = df.loc[:, ~df.columns.duplicated()]
        
        return df
    except Exception as e:
        logging.error(f"Error in write_dicts_to_excel function: {e}")

